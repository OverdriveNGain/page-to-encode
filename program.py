import pytesseract
from PIL import Image
from docx import Document
import os
import cv2
import numpy as np;
# from pynput import keyboard
import time
import keyboard
# import msvcrt

pytesseract.pytesseract.tesseract_cmd = 'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'
Attempt_to_combine_lines = True

def parse_true_or_false(s):
    return s.lower() in ['true', 'yes', 'y', 't']

print()
for line in open('options.txt', 'r').readlines():
    if not '=' in line:
        continue
    tokens = line.split("=")
    if tokens[0].strip() == 'Attempt_to_combine_lines':
        Attempt_to_combine_lines = parse_true_or_false(tokens[1].strip())
        print(f"Option detected: Attempt_to_combine_lines set to {Attempt_to_combine_lines}")
print()

def save_string_to_word_file(text, filename, image_path):
    # Create a new Word document
    document = Document()
    # Add the string as a paragraph to the document
    document.add_paragraph(text)
    document.add_picture(image_path)
    # Save the document to the specified file
    document.save(filename)

def extract_text_from_image(image_path):
    # Open the image file using PIL
    with Image.open(image_path) as img:
        # Use pytesseract to extract the text from the image
        text = pytesseract.image_to_string(img)
    return text

def save_string_to_file(text, filename):
    with open(filename, 'w') as file:
        file.write(text)

def list_files_in_directory(directory):
    # Get a list of all files in the directory
    files = os.listdir(directory)

    # Return only the files (not directories) in the list
    return [file for file in files if os.path.isfile(os.path.join(directory, file))]

def remove_file_extension(filename):
    # Split the filename into the base name and extension
    base_name, extension = os.path.splitext(filename)
    # Return just the base name
    return base_name

def preprocess_image_for_ocr(image_path, output_image_path):
    # Load the image
    img = cv2.imread(image_path)
    # Resize the image to the appropriate size
    img = cv2.resize(img, (0,0), fx=2, fy=2)
    # Convert the image to grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # Apply binarization and thresholding
    ret, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    # Apply noise reduction
    img_processed = cv2.fastNlMeansDenoising(thresh, None, 10, 7, 21)
    cv2.imwrite(output_image_path, gray)

def process_files():
    fileList = list_files_in_directory('input')
    for file in fileList:
        input_file_path = f'input/{file}'
        process_file_path = f'process/{file}'
        output_file_path = f'output/{remove_file_extension(file)}.docx'
        
        # Preprocess image
        # flatten_book_page(input_file_path, process_file_path) # doesnt work
        preprocess_image_for_ocr(input_file_path, process_file_path)
        # Extract text
        text = extract_text_from_image(input_file_path)
        # Remove possible unnecessary newlines
        if (Attempt_to_combine_lines):
            textlines = text.split("\n")
            i = 0
            while (i < len(textlines) - 1):
                if (i == len(textlines) - 1):
                    break

                if (len(textlines[i]) > 0 and len(textlines[i + 1]) > 0 and textlines[i][-1].islower() and textlines[i+1][0].islower()):
                    textlines[i] = textlines[i] + ' ' + textlines[i+1]
                    del textlines[i+1]
                
                i += 1
            text = "\n".join(textlines)
        # Save to word file
        save_string_to_word_file(text, output_file_path, input_file_path)
        print(f"Saved to {output_file_path}")

def on_press(key):
    global listener
    if (key == 'a'):
        print("a is pressed!")
    elif (key == 'q'):
        listener.stop()

print("+--------------------+")
print("|   Text Extractor   |")
print("+--------------------+")
print()
print('[1] Extract all files')
print('[2] Quit')

while True:
    keyPressed = input('\n>> ')
    print()
    if keyPressed == "1":
        process_files()
    if keyPressed == "2":
        break